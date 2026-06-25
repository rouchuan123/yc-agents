from docx import Document

from yc_agents.rag import loaders
from yc_agents.rag.loaders import load_docx, load_markdown, load_pdf


def test_load_markdown_preserves_source_and_text(tmp_path):
    path = tmp_path / "notes.md"
    path.write_text("# Background\n\nLocal skill-driven runtime.", encoding="utf-8")

    document = load_markdown(path)

    assert document["source"] == str(path)
    assert "skill-driven" in document["text"]
    assert document["metadata"]["file_name"] == "notes.md"
    assert document["metadata"]["file_type"] == "markdown"


def test_load_docx_reads_paragraphs(tmp_path):
    path = tmp_path / "notes.docx"
    doc = Document()
    doc.add_paragraph("Project background")
    doc.add_paragraph("Technical approach")
    doc.save(path)

    loaded = load_docx(path)

    assert "Project background" in loaded["text"]
    assert "Technical approach" in loaded["text"]
    assert loaded["metadata"]["file_type"] == "docx"


def test_load_pdf_skips_empty_pages(tmp_path, monkeypatch):
    path = tmp_path / "notes.pdf"
    path.write_bytes(b"%PDF fake")

    class FakePage:
        def __init__(self, text):
            self.text = text

        def extract_text(self):
            return self.text

    class FakeReader:
        def __init__(self, _path):
            self.pages = [FakePage("Project background"), FakePage(""), FakePage("Technical approach")]

    monkeypatch.setattr(loaders, "PdfReader", FakeReader)

    loaded = load_pdf(path)

    assert loaded["source"] == str(path)
    assert loaded["text"] == "Project background\nTechnical approach"
    assert loaded["metadata"]["file_type"] == "pdf"
