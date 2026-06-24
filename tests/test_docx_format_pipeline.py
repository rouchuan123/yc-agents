import json
import tempfile
import unittest
from pathlib import Path

from docx import Document

from yc_agents.docx_format.pipeline import normalize_docx


class TestDocxFormatPipeline(unittest.TestCase):
    def test_normalize_docx_writes_docx_and_audit_files(self):
        with tempfile.TemporaryDirectory() as tmp_dir:
            root = Path(tmp_dir)
            source = root / "messy.docx"
            output_dir = root / "out"
            document = Document()
            document.add_heading("Project Overview", level=1)
            document.add_paragraph("Body paragraph.")
            document.save(source)

            result = normalize_docx(
                source_path=source,
                output_dir=output_dir,
                template_name="report-standard",
                output_name="normalized",
            )

            self.assertTrue(result.output_docx.exists())
            self.assertTrue(result.audit_report.exists())
            self.assertTrue(result.audit_json.exists())
            data = json.loads(result.audit_json.read_text(encoding="utf-8"))
            self.assertIn(data["status"], ["passed", "passed_with_warnings"])

    def test_normalize_rejects_missing_source(self):
        with tempfile.TemporaryDirectory() as tmp_dir:
            with self.assertRaises(FileNotFoundError):
                normalize_docx(
                    source_path=Path(tmp_dir) / "missing.docx",
                    output_dir=Path(tmp_dir) / "out",
                )


if __name__ == "__main__":
    unittest.main()
