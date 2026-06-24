import tempfile
import unittest
from pathlib import Path

from docx import Document

from yc_agents.docx_format.analyzer import analyze_docx


class TestDocxAnalyzer(unittest.TestCase):
    def test_analyzes_headings_paragraphs_and_tables(self):
        with tempfile.TemporaryDirectory() as tmp_dir:
            path = Path(tmp_dir) / "draft.docx"
            document = Document()
            document.add_heading("Project Overview", level=1)
            document.add_paragraph("This is the first body paragraph.")
            table = document.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "Name"
            table.cell(0, 1).text = "Value"
            table.cell(1, 0).text = "Capacity"
            table.cell(1, 1).text = "100kW"
            document.save(path)

            model = analyze_docx(path)

            self.assertEqual(model.source_path, str(path))
            self.assertEqual(
                [block.text for block in model.blocks_by_type("heading")],
                ["Project Overview"],
            )
            self.assertEqual(
                [block.text for block in model.blocks_by_type("paragraph")],
                ["This is the first body paragraph."],
            )
            self.assertEqual(
                model.blocks_by_type("table")[0].rows[1],
                ["Capacity", "100kW"],
            )

    def test_detects_caption_like_paragraph(self):
        with tempfile.TemporaryDirectory() as tmp_dir:
            path = Path(tmp_dir) / "caption.docx"
            document = Document()
            document.add_paragraph("Figure 1-1 Site photo")
            document.save(path)

            model = analyze_docx(path)

            self.assertEqual(model.blocks[0].type, "caption")

    def test_rejects_non_docx(self):
        with tempfile.TemporaryDirectory() as tmp_dir:
            path = Path(tmp_dir) / "draft.txt"
            path.write_text("not docx", encoding="utf-8")

            with self.assertRaises(ValueError):
                analyze_docx(path)


if __name__ == "__main__":
    unittest.main()
