import tempfile
import unittest
from pathlib import Path

from docx import Document

from yc_agents.tools.docx_format_normalizer import DocxFormatNormalizerTool


class TestDocumentFormatNormalizerDemo(unittest.TestCase):
    def test_demo_flow_creates_reviewable_artifacts(self):
        with tempfile.TemporaryDirectory() as tmp_dir:
            workspace = Path(tmp_dir)
            source = workspace / "messy-demo.docx"
            document = Document()
            document.add_paragraph("Project Overview")
            document.add_paragraph("This draft has content but inconsistent formatting.")
            document.save(source)

            result = DocxFormatNormalizerTool(workspace).run(
                source_file="messy-demo.docx",
                output_name="demo-normalized",
            )

            self.assertTrue(result["ok"])
            self.assertTrue(result["output_docx"].endswith("demo-normalized.docx"))
            self.assertTrue(result["audit_report"].endswith("demo-normalized.audit.md"))


if __name__ == "__main__":
    unittest.main()
