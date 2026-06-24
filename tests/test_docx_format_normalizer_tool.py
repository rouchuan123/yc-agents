import tempfile
import unittest
from pathlib import Path

from docx import Document

from yc_agents.tools.docx_format_normalizer import DocxFormatNormalizerTool


class TestDocxFormatNormalizerTool(unittest.TestCase):
    def test_tool_normalizes_workspace_docx(self):
        with tempfile.TemporaryDirectory() as tmp_dir:
            workspace = Path(tmp_dir)
            source = workspace / "draft.docx"
            document = Document()
            document.add_heading("Project Overview", level=1)
            document.add_paragraph("Body paragraph.")
            document.save(source)

            result = DocxFormatNormalizerTool(workspace).run(
                source_file="draft.docx",
                template_name="report-standard",
                output_name="result",
            )

            self.assertTrue(result["ok"])
            self.assertEqual(result["template"], "report-standard")
            self.assertTrue((workspace / result["output_docx"]).exists())
            self.assertTrue((workspace / result["audit_report"]).exists())

    def test_tool_rejects_path_escape(self):
        with tempfile.TemporaryDirectory() as tmp_dir:
            workspace = Path(tmp_dir)

            with self.assertRaises(PermissionError):
                DocxFormatNormalizerTool(workspace).run(source_file="../draft.docx")


if __name__ == "__main__":
    unittest.main()
