import tempfile
import unittest
from pathlib import Path

from docx import Document

from yc_agents.docx_format.formatter import format_docx
from yc_agents.docx_format.models import DocumentBlock, DocumentModel
from yc_agents.docx_format.template import load_builtin_template


class TestDocxFormatter(unittest.TestCase):
    def test_formats_headings_paragraphs_tables_and_toc(self):
        with tempfile.TemporaryDirectory() as tmp_dir:
            output_path = Path(tmp_dir) / "normalized.docx"
            model = DocumentModel(
                source_path="draft.docx",
                blocks=[
                    DocumentBlock(id="b1", type="heading", text="Project Overview", level=1),
                    DocumentBlock(id="b2", type="paragraph", text="Body paragraph."),
                    DocumentBlock(id="b3", type="caption", text="Table 1-1 Parameters"),
                    DocumentBlock(
                        id="b4",
                        type="table",
                        rows=[["Name", "Value"], ["Capacity", "100kW"]],
                    ),
                ],
            )

            format_docx(model, load_builtin_template("report-standard"), output_path)

            self.assertTrue(output_path.exists())
            document = Document(output_path)
            texts = [paragraph.text for paragraph in document.paragraphs]
            self.assertIn("Project Overview", texts)
            self.assertIn("Body paragraph.", texts)
            self.assertIn("Table 1-1 Parameters", texts)
            self.assertEqual(document.tables[0].cell(1, 1).text, "100kW")
            self.assertIn("Table of Contents", "\n".join(texts))

    def test_rejects_empty_output_path(self):
        model = DocumentModel(source_path="draft.docx")

        with self.assertRaises(ValueError):
            format_docx(model, load_builtin_template("report-standard"), "")


if __name__ == "__main__":
    unittest.main()
