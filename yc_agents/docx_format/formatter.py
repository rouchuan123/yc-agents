from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ALIGNMENTS = {
    "center": WD_PARAGRAPH_ALIGNMENT.CENTER,
    "left": WD_PARAGRAPH_ALIGNMENT.LEFT,
    "right": WD_PARAGRAPH_ALIGNMENT.RIGHT,
    "justify": WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
}


def format_docx(document_model, template_rules, output_path):
    output = Path(output_path) if output_path else None
    if output is None:
        raise ValueError("output_path is required")

    output.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    _apply_page_setup(doc, template_rules)
    _configure_normal_style(doc, template_rules)
    _add_toc(doc, template_rules)
    _add_content_section(doc, template_rules)

    for block in document_model.blocks:
        if block.type == "heading":
            _add_heading(doc, block, template_rules)
        elif block.type == "paragraph":
            _add_body_paragraph(doc, block.text, template_rules)
        elif block.type == "caption":
            _add_caption(doc, block.text, template_rules)
        elif block.type == "table":
            _add_table(doc, block.rows, template_rules)
        elif block.type == "image" and block.image_path:
            _add_image(doc, block.image_path)

    _enable_update_fields_on_open(doc)
    doc.save(output)
    return output


def _apply_page_setup(doc, template_rules):
    margins = template_rules.page["margins_cm"]
    for section in doc.sections:
        section.top_margin = Cm(margins["top"])
        section.bottom_margin = Cm(margins["bottom"])
        section.left_margin = Cm(margins["left"])
        section.right_margin = Cm(margins["right"])
        section.header_distance = Cm(template_rules.page.get("header_distance_cm", 1.5))
        section.footer_distance = Cm(template_rules.page.get("footer_distance_cm", 1.75))


def _configure_normal_style(doc, template_rules):
    body = template_rules.styles["body"]
    style = doc.styles["Normal"]
    style.font.name = body["east_asia_font"]
    style._element.rPr.rFonts.set(qn("w:eastAsia"), body["east_asia_font"])
    style.font.size = Pt(body["font_size_pt"])
    style.paragraph_format.space_before = Pt(body["space_before_pt"])
    style.paragraph_format.space_after = Pt(body["space_after_pt"])
    style.paragraph_format.line_spacing = Pt(body["line_spacing_pt"])
    style.paragraph_format.first_line_indent = Pt(body["first_line_indent_pt"])


def _add_toc(doc, template_rules):
    if not template_rules.table_of_contents.get("enabled"):
        return
    heading = doc.add_paragraph()
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    heading.add_run("Table of Contents").bold = True
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    run._r.append(begin)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = template_rules.table_of_contents["field_code"]
    run._r.append(instr)
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    run._r.append(separate)
    run.add_text("Update fields in Word or WPS.")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(end)


def _add_content_section(doc, template_rules):
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    margins = template_rules.page["margins_cm"]
    section.top_margin = Cm(margins["top"])
    section.bottom_margin = Cm(margins["bottom"])
    section.left_margin = Cm(margins["left"])
    section.right_margin = Cm(margins["right"])
    if template_rules.page_numbers.get("enabled"):
        paragraph = section.footer.paragraphs[0]
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = paragraph.add_run()
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        run._r.append(begin)
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = "PAGE \\* MERGEFORMAT"
        run._r.append(instr)
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        run._r.append(end)


def _add_heading(doc, block, template_rules):
    level = min(block.level or 1, 3)
    paragraph = doc.add_heading(block.text, level=level)
    style_rules = template_rules.styles[f"heading_{level}"]
    paragraph.alignment = ALIGNMENTS[style_rules["alignment"]]
    for run in paragraph.runs:
        run.font.name = style_rules["east_asia_font"]
        run._element.rPr.rFonts.set(qn("w:eastAsia"), style_rules["east_asia_font"])
        run.font.size = Pt(style_rules["font_size_pt"])
        run.font.bold = style_rules["bold"]


def _add_body_paragraph(doc, text, template_rules):
    paragraph = doc.add_paragraph(text)
    body = template_rules.styles["body"]
    paragraph.alignment = ALIGNMENTS[body["alignment"]]
    paragraph.paragraph_format.line_spacing = Pt(body["line_spacing_pt"])
    paragraph.paragraph_format.first_line_indent = Pt(body["first_line_indent_pt"])
    for run in paragraph.runs:
        run.font.name = body["east_asia_font"]
        run._element.rPr.rFonts.set(qn("w:eastAsia"), body["east_asia_font"])
        run.font.size = Pt(body["font_size_pt"])


def _add_caption(doc, text, template_rules):
    paragraph = doc.add_paragraph(text)
    caption = template_rules.styles["caption"]
    paragraph.alignment = ALIGNMENTS[caption["alignment"]]
    paragraph.paragraph_format.space_before = Pt(caption["space_before_pt"])
    paragraph.paragraph_format.space_after = Pt(caption["space_after_pt"])
    for run in paragraph.runs:
        run.font.name = caption["east_asia_font"]
        run._element.rPr.rFonts.set(qn("w:eastAsia"), caption["east_asia_font"])
        run.font.size = Pt(caption["font_size_pt"])


def _add_table(doc, rows, template_rules):
    if not rows:
        return
    cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    for row_index, row in enumerate(rows):
        for col_index in range(cols):
            table.cell(row_index, col_index).text = (
                row[col_index] if col_index < len(row) else ""
            )


def _add_image(doc, image_path):
    path = Path(image_path)
    if not path.exists():
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    paragraph.add_run().add_picture(str(path), width=Cm(15))


def _enable_update_fields_on_open(doc):
    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")
