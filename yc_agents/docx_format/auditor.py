from pathlib import Path

from docx import Document

from yc_agents.docx_format.models import AuditCheck, AuditReport


def audit_docx(output_path, template_rules, source_model=None):
    path = Path(output_path)
    checks = []
    if path.exists() and path.stat().st_size > 0:
        checks.append(AuditCheck("file_exists", "passed", "Output DOCX exists."))
    else:
        checks.append(AuditCheck("file_exists", "failed", "Output DOCX is missing or empty."))
        return AuditReport("failed", str(path), checks)

    document = Document(path)
    checks.extend(_check_page_margins(document, template_rules))
    checks.extend(_check_normal_style(document, template_rules))
    checks.append(_check_toc(document, template_rules))
    checks.append(_check_page_numbers(document, template_rules))

    if source_model and source_model.unsupported_objects:
        count = len(source_model.unsupported_objects)
        checks.append(
            AuditCheck(
                "unsupported_objects",
                "warning",
                f"Detected {count} unsupported source object(s); review manually.",
            )
        )

    status = "passed"
    if any(check.status == "failed" for check in checks):
        status = "failed"
    elif any(check.status == "warning" for check in checks):
        status = "passed_with_warnings"
    return AuditReport(status, str(path), checks)


def _check_page_margins(document, template_rules):
    expected = template_rules.page["margins_cm"]
    section = document.sections[0]
    actual = {
        "top": round(section.top_margin.cm, 2),
        "bottom": round(section.bottom_margin.cm, 2),
        "left": round(section.left_margin.cm, 2),
        "right": round(section.right_margin.cm, 2),
    }
    if actual == expected:
        return [AuditCheck("page_margins", "passed", "Page margins match template.")]
    return [
        AuditCheck(
            "page_margins",
            "failed",
            f"Expected margins {expected}, got {actual}.",
        )
    ]


def _check_normal_style(document, template_rules):
    expected_size = template_rules.styles["body"]["font_size_pt"]
    style = document.styles["Normal"]
    actual_size = round(style.font.size.pt, 2) if style.font.size else None
    if actual_size == expected_size:
        return [AuditCheck("normal_style", "passed", "Normal style font size matches.")]
    return [
        AuditCheck(
            "normal_style",
            "failed",
            f"Expected Normal font size {expected_size}, got {actual_size}.",
        )
    ]


def _check_toc(document, template_rules):
    if not template_rules.table_of_contents.get("enabled"):
        return AuditCheck("table_of_contents", "passed", "TOC not required.")
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    if "Table of Contents" in text:
        return AuditCheck("table_of_contents", "passed", "TOC field marker exists.")
    return AuditCheck("table_of_contents", "failed", "TOC field marker missing.")


def _check_page_numbers(document, template_rules):
    if not template_rules.page_numbers.get("enabled"):
        return AuditCheck("page_numbers", "passed", "Page numbers not required.")
    for section in document.sections:
        footer_text = "\n".join(paragraph.text for paragraph in section.footer.paragraphs)
        if footer_text or section.footer.paragraphs:
            return AuditCheck("page_numbers", "passed", "Footer exists for page numbers.")
    return AuditCheck("page_numbers", "failed", "Footer page-number area missing.")
