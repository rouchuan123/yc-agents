import re
import zipfile
from pathlib import Path

from docx import Document

from yc_agents.docx_format.models import DocumentBlock, DocumentModel, UnsupportedObject


CAPTION_RE = re.compile(
    r"^(Figure|Fig\.|Table|\u56fe|\u8868)\s*[\d\u4e00-\u5341]+[-.\uff0d-]?\d*",
    re.IGNORECASE,
)


def analyze_docx(file_path, media_output_dir=None):
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"DOCX file not found: {path}")
    if path.suffix.lower() != ".docx":
        raise ValueError(f"Expected a .docx file, got: {path}")

    document = Document(path)
    blocks = []
    block_index = 1

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        block_type, level = _classify_paragraph(paragraph, text)
        blocks.append(
            DocumentBlock(
                id=f"block_{block_index:04d}",
                type=block_type,
                text=text,
                level=level,
                style_name=paragraph.style.name if paragraph.style else None,
                format=_paragraph_format(paragraph),
            )
        )
        block_index += 1

    for table in document.tables:
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        blocks.append(
            DocumentBlock(
                id=f"block_{block_index:04d}",
                type="table",
                rows=rows,
                style_name=table.style.name if table.style else None,
            )
        )
        block_index += 1

    media_dir = None
    image_blocks = []
    if media_output_dir is not None:
        media_dir = _extract_media(path, Path(media_output_dir), block_index)
        for image_index, image_path in enumerate(
            sorted(Path(media_dir).glob("*")),
            start=block_index,
        ):
            image_blocks.append(
                DocumentBlock(
                    id=f"block_{image_index:04d}",
                    type="image",
                    image_path=str(image_path),
                )
            )
    blocks.extend(image_blocks)

    return DocumentModel(
        source_path=str(path),
        blocks=blocks,
        media_dir=media_dir,
        unsupported_objects=_detect_unsupported_objects(path),
    )


def _classify_paragraph(paragraph, text):
    style_name = paragraph.style.name if paragraph.style else ""
    if style_name.startswith("Heading"):
        match = re.search(r"(\d+)$", style_name)
        level = int(match.group(1)) if match else 1
        return "heading", min(level, 3)
    if CAPTION_RE.match(text):
        return "caption", None
    return "paragraph", None


def _paragraph_format(paragraph):
    first_run = paragraph.runs[0] if paragraph.runs else None
    font = first_run.font if first_run else None
    return {
        "alignment": str(paragraph.alignment) if paragraph.alignment is not None else None,
        "style_name": paragraph.style.name if paragraph.style else None,
        "font": font.name if font and font.name else None,
        "font_size_pt": font.size.pt if font and font.size else None,
        "bold": font.bold if font else None,
        "italic": font.italic if font else None,
    }


def _extract_media(docx_path, media_output_dir, start_index):
    media_output_dir.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(docx_path, "r") as archive:
        media_names = [
            name
            for name in archive.namelist()
            if name.startswith("word/media/") and not name.endswith("/")
        ]
        for offset, name in enumerate(media_names):
            suffix = Path(name).suffix or ".bin"
            target = media_output_dir / f"image_{start_index + offset:04d}{suffix}"
            target.write_bytes(archive.read(name))
    return str(media_output_dir)


def _detect_unsupported_objects(docx_path):
    unsupported = []
    with zipfile.ZipFile(docx_path, "r") as archive:
        names = archive.namelist()
        markers = {
            "word/comments.xml": "comments",
            "word/charts/": "chart",
            "word/embeddings/": "embedded_object",
            "word/vbaProject.bin": "macro",
        }
        for marker, object_type in markers.items():
            if any(name == marker or name.startswith(marker) for name in names):
                unsupported.append(
                    UnsupportedObject(type=object_type, location=marker)
                )
    return unsupported
